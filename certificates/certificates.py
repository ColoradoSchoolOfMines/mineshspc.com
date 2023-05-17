import csv
import os
import time
import shutil
import subprocess
from docx import Document
from docx.shared import Pt

file_name = input("Enter a csv to read: ")
start = time.time()
count = 0

with open(file_name, 'r') as file:
    for line in csv.reader(file):
        student = line[0]

        name = f'{student}-HSPC-Certificate.docx'
        shutil.copyfile('HSPC-Certificate.docx', name)

        document = Document(name)
        document.paragraphs[4].text = student
        document.paragraphs[4].runs[0].font.size = Pt(48)
        document.paragraphs[4].runs[0].font.name = 'Plantagenet Cherokee'
        document.save(name)

        # convert .docx -> .pdf; libreoffice + libreofficewriter need to be installed
        subprocess.call(['soffice', '--headless', '--convert-to', 'pdf:writer_pdf_Export', '--outdir', "advanced-certificates/", name])

        os.remove(name)
        count = count + 1

end = time.time()
print(f"Generated {count} certificates in {end - start:.1f}s")
